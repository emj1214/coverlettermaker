# Author: Emma Johnson
# Date Last Updated: December 6, 2022
# URL: https://github.com/emj1214/coverlettermaker

import enquiries
import docx
from docx.shared import Inches
import datetime
import docx2pdf
import hyperlink

# This last import should not go into GitHub – it contains personal information such as phone numbers, emails, etc. 
# A template has been created so that you can edit your own. It should not take long to fill out and saves you time by aut-filling multiple variables.
# change this import to the name of YOUR privatevars file if you change it from privatevarsTEMPLATE
import privatevars 

# Create the word doc
doc = docx.Document()

# Brief welcome before the code asks its questions
print("\n\n\n\n\n\n\nWelcome! Let's begin!")
print("\n\n------------\n\n")

# Set Yes and No options for future selectors
yn = ["Yes", "No"]

# Who are we addressing the letter to?
## Gather the first name, if applicable
while True:
    firstname = input("What is the hiring manager's FIRST name? \nLeave blank if unknown or if addressing to the company's Hiring Team. \n\n").lstrip(" ").rstrip(" ").capitalize()
    ## Determine if a space after the firstname is necessary (no space is needed if there is no first name) 
    if firstname != "":
        firstnamespace = " "
    else:
        firstnamespace = ""
    correct = enquiries.choose(f"\nIs '{firstname}' correct so far?", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

## Gather the last name, if applicable
while True:
    lastname = input("What is the hiring manager's LAST name, OR, if addressing to a HIRING TEAM, what is the team called? \nLeave blank if unknown. \n\n").lstrip(" ").rstrip(" ").capitalize()
    ## Determine if a space after the last name is necessary (no space is needed if there is no last name) 
    if lastname != "":
        lastnamespace = " "
    else:
        lastnamespace = ""
    correct = enquiries.choose(f"\nIs '{firstname + firstnamespace + lastname}' correct so far?", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

## Gather the prefix, if applicable
while True:
    prefix = input("What is the hiring manager's prefix (Ms/Mrs/Mr/Mx/etc)? \nLeave blank if unknown or if addressing to the company's Hiring Team. \n\n").lstrip(" ").rstrip(" ").capitalize()
    ## Determine if a space after the prefix is necessary (no space is needed if there is no prefix) 
    if prefix != "":
        prefixspace = " "
        ## Add a period to an abbreviation, in case the user forgot
        if (prefix[len(prefix)-1] != ".") and (len(prefix)<4):
            prefix = prefix + "."
    else:
        prefixspace = ""
    correct = enquiries.choose(f"\nIs '{prefix + firstnamespace + firstname + lastnamespace + lastname}' correct so far?", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

## Gather mailing address
while True:
    while True:
        try:
            address = input("What is the hiring manager's mailing address? Insert a $ between the street address and the city \n(ex: 1234 Avenue Street $ New York, NY) \n\n").lstrip(" ").rstrip(" ")
            address = address.split(sep="$")
            address[0] = address[0].lstrip().rstrip().title()
            address[1] = address[1].lstrip().rstrip()
            break
        except IndexError as e:
            print("Oops! Address needs both lines to continue. Try it again. \n")
    correct = enquiries.choose(f"\nIs this correct: \n{address[0]} \n{address[1]}", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

# My personal info as a default from file 'privatevars.py', option to replace info if I decide to share with a friend :)
myname = enquiries.choose(f'Are you {privatevars.myfirstname + " " + privatevars.mylastname}?', yn)
if myname == "Yes":
    myfirstname = privatevars.myfirstname
    mylastname = privatevars.mylastname
    myprefix = privatevars.myprefix
    whodesc = privatevars.whodesc
    myphone = privatevars.myphone
    myemail = privatevars.myemail
    myaddress = privatevars.myaddress
    mypronouns = privatevars.mypronouns

else:
    print("\n\n------------\n\n")

    while True:
        myfirstname = input("What is your FIRST name? \n\n").lstrip(" ").rstrip(" ").capitalize()
        print("\n" + myfirstname)
        print("\n\n------------\n\n")

        mylastname = input("What is your LAST name? \n\n").lstrip(" ").rstrip(" ").capitalize()
        print("\n" + mylastname)
        print("\n\n------------\n\n")

        myprefix = input("What prefix do you want to use? (Ms/Mrs/Mr/Mx/etc) \n\n").lstrip(" ").rstrip(" ").capitalize()
        ## Add a period to an abbreviation, in case the user forgot
        if (len(myprefix)>0) and (myprefix[len(myprefix)-1] != ".") and (len(myprefix)<4):
            myprefix = myprefix + "."
        print("\n" + myprefix)
        print("\n\n------------\n\n")

        whodesc = input("Describe yourself: I am a...(ex: third-year Dartmouth student) \n\n").lstrip(" ").rstrip(" ")
        print("\n" + whodesc)
        print("\n\n------------\n\n")

        myphone = input("What's your phone number? \n\n").lstrip(" ").rstrip(" ")
        print("\n" + myphone)
        print("\n\n------------\n\n")

        myemail = input("What's your email? \n\n").lstrip(" ").rstrip(" ")
        print("\n" + myemail)
        print("\n\n------------\n\n")

        mypronouns = input("What are your pronouns? \n\n").lstrip(" ").rstrip(" ")
        print("\n" + mypronouns)
        print("\n\n------------\n\n")

        correct = enquiries.choose(f"\nIs the following correct: \n{myfirstname} \n{mylastname} \n{myprefix} \n{whodesc} \n{myphone} \n{myemail} \n{mypronouns}", yn)
        if correct == "Yes":
            print("\n\n------------\n\n")
            break

# What's the position title?
while True:
    while True: 
        try:
            position = input("What's the title of the position you're applying for? (ex: Software Development Intern) \n\n").lstrip(" ").rstrip(" ").title()
            # Determine if in a sentence it is more appropriate to use 'a' or 'an' before the job title
            vowels = ['a', 'e', 'i', 'o', 'u', 'A', 'E', 'I', 'O', 'U']
            if position[0] in vowels:
                positionart = "an " + position
            else :
                positionart = "a " + position
            break
        except IndexError as e:
            print("Oops! Position cannot be skipped. Try it again. \n")
    correct = enquiries.choose(f"\nIs this correct: \n{position}", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

# Gathering the company name
while True:
    company = input("What is the company called? \n\n")
    correct = enquiries.choose(f"\nIs this correct: \n{company}", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

# A drop-down selector of values given in privatevars.py
while True:
    try:
        values = enquiries.choose("What THREE values do you want to highlight? Use the space bar and arrow keys to select items.", privatevars.valopts, multi=True)
        value1 = values[0]
        value2 = values[1]
        value3 = values[2]
        break
    except IndexError as e:
        print("Oops! You didn't select THREE values. Try it again.")
print("\n\n------------\n\n")

# What sort of challenges might the team be facing that I can contribute towards?
# This is usually included in the job description
while True:
    teamchallenge = input("From the job description, what's the team's goal that you'll help with? (ex: create an efficient and community-oriented work-environment) \n\n In other words, finish this statement: \n 'I will help you and your team...'")
    correct = enquiries.choose(f"\nIs this correct: \n{teamchallenge}", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

# I have a short few sentences I add about my Clifton Strengths from Gallup that I like to include when applicable (i.e. the hiring manager has their Strengths on LinkedIn)
usestrengths = enquiries.choose("Do you want to include your Clifton Strengths?", yn)
if usestrengths == "Yes":
    strengths = privatevars.strengthssent
else:
    strengths = ""
print("\n\n------------\n\n")

# I have a blurb about how I like to learn and what I want to discover in the job I'm applying for.
uselearning = enquiries.choose("Do you want to include your sentence on learning in the workplace?", yn)
if uselearning == "Yes":
    focus = input("What topic will you learn about through this job? (ex: data analytics) \n\n")
    learningsent = privatevars.learningsentp1 + focus + privatevars.learningsentp2 + company + privatevars.learningsentp3
else:
    learningsent = None
print("\n\n------------\n\n")

# Ask where I found the poisition
while True:
    found = input("How/Where did you find this position? (ex: on LinkedIn) \n\n")
    correct = enquiries.choose(f"\nIs this correct: \n{found}", yn)
    if correct == "Yes":
        print("\n\n------------\n\n")
        break

# --- Begin formatting document ---
# Set bottom margin to half an inch to fit my signature
doc.sections[0].bottom_margin = Inches(0.5)

# Add today's date to the heading
today = datetime.date.today()
todaystr = today.strftime("%d %B %Y")
date = doc.add_paragraph(f'{todaystr}')

# Extra space (empty paragraph)
doc.add_paragraph()

# Fill out letter address 
managerinfo = doc.add_paragraph(f"{prefix} {firstname} {lastname} \n{company} \n{address[0]} \n{address[1]}")

# Extra space (empty paragraph)
doc.add_paragraph()

# Fill in greeting
# Note: you can see here why prefixspace is necessary: if I addressed the letter to the hiring team without a prefix, there would be a double space before the team's title I would have to manually correct.
greeting = doc.add_paragraph(f"Dear {prefix}{prefixspace}{lastname},")

# Fill in introduction paragraph
introduction = doc.add_paragraph(
    f"My name is {myfirstname} {mylastname}, and I'm {whodesc} excited to be submitting myself for the {position} opportunity with {company}."
)

# Fill in the first body paragraph about why I would be a good fit
body1 = doc.add_paragraph(
    f"As {positionart}, I believe I would make a valuable asset to your team. My {value1}, {value2}, and {value3} will help you and your team {teamchallenge}. {strengths}"
)

# Fill in the second body paragraph about why the position would benefit me
body2 = doc.add_paragraph(
    f"The job description {found} that led me to your company highlighted several of my own values as those emphasized in the workplace, and as a part of {company}, it would be natural for me to exercise my skills with your team. {learningsent}" 
)

# Ask if the user wants to use the hard skills bullet points, and if so, which they'd like to use
hardskills = enquiries.choose("Do you want to include your section on hard skills?", yn)
if hardskills == "Yes":
    # provides hard skills as options to select
    while True:
        try:
            achieve = enquiries.choose("Which THREE hard skills?", privatevars.achieveoptslist, multi=True)
            links = []

            # for each selected hard skill, the respective URL to see the project associated with the skill is added to our temporary links list
            for i in achieve:
                index = privatevars.achieveoptslist.index(achieve[achieve.index(i)])
                links.append(privatevars.mylinks[index])

            # Fill in opening for the hard skills / my achievements paragraph
            doc.add_paragraph("My hard skills (and the projects for which I used them) that you may find relevant to your team include: ")

            # Create the first bullet point by writing the first half of the sentence, the hyperlinked text using the appropriate URL from the links list, and the second half of the sentence if applicable
            bullet1 = doc.add_paragraph(style='List Bullet')
            bullet1.add_run(f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[0])][0]}')
            hyperlink.add_hyperlink(paragraph=bullet1,text=f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[0])][1]}', url=f'{links[0]}')
            bullet1.add_run(f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[0])][2]}')

            # Create the second bullet point by writing the first half of the sentence, the hyperlinked text using the appropriate URL from the links list, and the second half of the sentence if applicable
            bullet2 = doc.add_paragraph(style='List Bullet')
            bullet2.add_run(f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[1])][0]}')
            hyperlink.add_hyperlink(paragraph=bullet2,text=f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[1])][1]}', url=f'{links[1]}')
            bullet2.add_run(f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[1])][2]}')

            # Create the third bullet point by writing the first half of the sentence, the hyperlinked text using the appropriate URL from the links list, and the second half of the sentence if applicable
            bullet3 = doc.add_paragraph(style='List Bullet')
            bullet3.add_run(f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[2])][0]}')
            hyperlink.add_hyperlink(paragraph=bullet3,text=f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[2])][1]}', url=f'{links[2]}')
            bullet3.add_run(f'{privatevars.achieveopts[privatevars.achieveoptslist.index(achieve[2])][2]}')

            break

        except IndexError as e:
            print("Oops! You didn't select THREE values. Try it again.")

# Fill in the conclusion paragraph with the user's contact info
conclusion = doc.add_paragraph(
    f"When it is most convenient for you, I would appreciate the opportunity to discuss where my skills could benefit your company. I can be reached at {myemail} or at {myphone}. Thank you so much for your consideration, and I look forward to hearing from you soon."
)

# Extra space (empty paragraph)
doc.add_paragraph()

# Signature
sincerely = doc.add_paragraph("Sincerely,")
signature = doc.add_paragraph(f"{myprefix} {myfirstname} {mylastname} \n{mypronouns} \n{myphone} \n{myemail}")

# Save doc to current directory
doc.save(f'{myfirstname}{mylastname}CoverLetter{"".join(company.split())}.docx')

print("\n\n------------\n\n")

# Heads up that the doc was successfully generated
print("A Word file has been generated.")
print("\n\n------------\n\n")

# The option to turn the .docx file into a .pdf file
pdf = enquiries.choose("Do you want a pdf now?", yn)
if pdf == "Yes":
    docx2pdf.convert(f'{myfirstname}{mylastname}CoverLetter{"".join(company.split())}.docx', f'{myfirstname}{mylastname}CoverLetter{"".join(company.split())}.pdf')